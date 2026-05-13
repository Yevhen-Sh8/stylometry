#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
core/monitoring_forms.py
========================
Формування формалізованого бланку Анкети добового моніторингу (Додаток 1)
відповідно до Методики щодо процедур моніторингу інформації у відкритих
джерелах та її обробки (затв. наказом НУЗРКС МОУ № 46 від 28.11.2022).

Структура бланку відтворює оригінальний Додаток 1 (зразок СЦМОУ від
18.10.2023): горизонтальна таблиця з рядком на кожне опрацьоване джерело,
секційними заголовками за сферою та підписним блоком.

Колонки:
  №  |  Необхідність реагування (так/ні)  |  Пріоритет  |
  Основні відомості  |  Пропозиції щодо реагування / посилання
"""

from __future__ import annotations

from datetime import datetime
from io import BytesIO
from typing import Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ─────────────────────────────────────────────────────────────────────────────
#  КОНСТАНТИ ДИЗАЙНУ
#  Відтворюють оформлення оригінального бланку СЦМОУ.
# ─────────────────────────────────────────────────────────────────────────────

# Ширини колонок (сума = ~18.5 cm при полях 2.5 / 1.5)
_COL_WIDTHS = (Cm(1.0), Cm(2.5), Cm(2.2), Cm(8.3), Cm(4.5))

# Колір секційних заголовків (лосось, як у оригіналі)
_SECTION_BG = "FFC0BC"

# Колір рядка для критичних позицій (SSS)
_CRITICAL_BG = "FFD7D7"

# Розмір шрифту основного тексту
_FONT_SIZE_BODY = 10
_FONT_SIZE_HEAD = 11

# ─────────────────────────────────────────────────────────────────────────────
#  РЕКОМЕНДАЦІЇ ЗА ГРЕЙДОМ (Методика, розділ 1)
# ─────────────────────────────────────────────────────────────────────────────
_GRADE_RECOMMENDATIONS: dict[str, str] = {
    "F":   "Малий інтерес. Розміщення та зберігання відомостей у базі даних "
           "без опрацювання.",
    "B":   "Точковий інтерес. Зберігання у базі даних з обмеженим "
           "опрацюванням.",
    "S":   "Однозначний інтерес. Зберігання у базі даних з подальшим "
           "опрацюванням та надання пропозицій щодо реагування.",
    "SS":  "Вагомий інтерес. Повідомлення керівника підрозділу, зберігання "
           "у базі даних з якнайшвидшим опрацюванням (за потреби у неробочий "
           "час) та пропозиціями щодо реагування.",
    "SSS": "Критичний інтерес. Позачергове (негайне) опрацювання у неробочий "
           "час, повідомлення керівника підрозділу та інших осіб, що "
           "приймають рішення.",
}

_GRADE_LABELS: dict[str, str] = {
    "F":   "малий інтерес",
    "B":   "точковий",
    "S":   "однозначний",
    "SS":  "вагомий",
    "SSS": "критичний",
}


def _requires_response(grade: str) -> str:
    """ТАК / НІ відповідно до Методики (реагування — з S-грейду)."""
    return "ТАК" if grade in {"S", "SS", "SSS"} else "НІ"


# ─────────────────────────────────────────────────────────────────────────────
#  XML-ХЕЛПЕРИ ДЛЯ DOCX
# ─────────────────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str) -> None:
    """Заливка клітинки кольором (W:shd)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_borders(cell, **kwargs) -> None:
    """Встановлення меж клітинки (top/bottom/left/right)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        val = kwargs.get(side, "single")
        sz  = kwargs.get(f"{side}_sz", "4")
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   val)
        el.set(qn("w:sz"),    str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _para_run(cell,
              text: str,
              bold: bool = False,
              italic: bool = False,
              size_pt: int = _FONT_SIZE_BODY,
              align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
              underline: bool = False,
              add_newline: bool = False) -> None:
    """Додає параграф з текстом до клітинки (не замінює існуючий)."""
    para = cell.add_paragraph()
    para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.size = Pt(size_pt)
    if add_newline:
        cell.add_paragraph()


def _clear_and_write(cell,
                     text: str,
                     bold: bool = False,
                     italic: bool = False,
                     size_pt: int = _FONT_SIZE_BODY,
                     align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT) -> None:
    """Очищає клітинку та записує один параграф."""
    for p in cell.paragraphs:
        p.clear()
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)


def _merge_row(table, row_idx: int) -> None:
    """Об'єднує всі клітинки рядка (секційний заголовок)."""
    row = table.rows[row_idx]
    cells = row.cells
    cells[0].merge(cells[-1])


# ─────────────────────────────────────────────────────────────────────────────
#  ПОБУДОВА РЯДКІВ ТАБЛИЦІ
# ─────────────────────────────────────────────────────────────────────────────

def _add_section_header(table, title: str) -> None:
    """Додає рядок-секційний заголовок (об'єднані клітинки, лосось)."""
    row = table.add_row()
    row_idx = len(table.rows) - 1
    _merge_row(table, row_idx)
    cell = row.cells[0]
    _set_cell_bg(cell, _SECTION_BG)
    _clear_and_write(
        cell, title,
        bold=True,
        size_pt=_FONT_SIZE_HEAD,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _source_row_info(
    source: dict,
    flagged_labels: set[str],
    breakdown_map: dict[str, float],
) -> str:
    """Формує текст колонки «Основні відомості» для одного джерела."""
    parts: list[str] = []

    title = (source.get("display_title")
             or source.get("label")
             or source.get("original_name")
             or "—")
    domain = source.get("domain") or ""
    url    = source.get("url") or source.get("local_text_url") or ""
    label  = source.get("label") or source.get("alias") or ""

    # Назва джерела
    parts.append(title)

    # Домен / URL
    if domain and url:
        parts.append(f"{domain} — {url}")
    elif url:
        parts.append(url)
    elif domain:
        parts.append(domain)

    # Оцінка ризику джерела (I_source)
    i_src = breakdown_map.get(label)
    if i_src is not None:
        parts.append(f"Індикатор ризику джерела I_source = {i_src:.4f}")

    # Позначка підозрілої пари
    if label in flagged_labels:
        parts.append("⚠ Входить до підозрілої стилометричної пари.")

    return "\n".join(parts)


def _source_row_proposal(
    grade: str,
    source: dict,
    manifestation_label: str,
    custom_recommendation: str,
) -> str:
    """Формує текст колонки «Пропозиції» для одного джерела."""
    parts: list[str] = []

    if custom_recommendation:
        parts.append(custom_recommendation.strip())
    else:
        parts.append(_GRADE_RECOMMENDATIONS.get(grade, ""))

    if manifestation_label:
        parts.append(f"Вид прояву DIMs: {manifestation_label}.")

    url = source.get("url") or ""
    if url:
        parts.append(f"Посилання: {url}")

    return "\n".join(p for p in parts if p)


def _add_source_row(
    table,
    num: int,
    source: dict,
    grade: str,
    manifestation_label: str,
    custom_recommendation: str,
    flagged_labels: set[str],
    breakdown_map: dict[str, float],
) -> None:
    """Додає рядок таблиці для одного джерела."""
    row = table.add_row()
    cells = row.cells

    for idx, w in enumerate(_COL_WIDTHS):
        cells[idx].width = w
        cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.TOP

    if grade == "SSS":
        for cell in cells:
            _set_cell_bg(cell, _CRITICAL_BG)

    _clear_and_write(cells[0], str(num),
                     align=WD_ALIGN_PARAGRAPH.CENTER,
                     size_pt=_FONT_SIZE_BODY)

    resp = _requires_response(grade)
    _clear_and_write(cells[1], resp,
                     bold=(resp == "ТАК"),
                     align=WD_ALIGN_PARAGRAPH.CENTER,
                     size_pt=_FONT_SIZE_BODY)

    grade_label = _GRADE_LABELS.get(grade, "")
    grade_text = f"{grade}\n({grade_label})"
    _clear_and_write(cells[2], grade_text,
                     bold=True,
                     align=WD_ALIGN_PARAGRAPH.CENTER,
                     size_pt=_FONT_SIZE_BODY)

    info_text = _source_row_info(source, flagged_labels, breakdown_map)
    _clear_and_write(cells[3], info_text, size_pt=_FONT_SIZE_BODY)

    proposal_text = _source_row_proposal(
        grade, source, manifestation_label, custom_recommendation,
    )
    _clear_and_write(cells[4], proposal_text, size_pt=_FONT_SIZE_BODY)


def _add_dims_summary_row(
    table,
    num: int,
    r_dims: float,
    grade: str,
    indicators: dict,
    flagged_pairs: list[dict],
) -> None:
    """Додає зведений рядок DIMS-аналізу (стилометрія + індикатори)."""
    row = table.add_row()
    cells = row.cells
    for idx, w in enumerate(_COL_WIDTHS):
        cells[idx].width = w
        cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.TOP

    _clear_and_write(cells[0], str(num),
                     align=WD_ALIGN_PARAGRAPH.CENTER,
                     size_pt=_FONT_SIZE_BODY)
    resp = _requires_response(grade)
    _clear_and_write(cells[1], resp,
                     bold=(resp == "ТАК"),
                     align=WD_ALIGN_PARAGRAPH.CENTER,
                     size_pt=_FONT_SIZE_BODY)
    grade_text = f"{grade}\n({_GRADE_LABELS.get(grade, '')})"
    _clear_and_write(cells[2], grade_text,
                     bold=True,
                     align=WD_ALIGN_PARAGRAPH.CENTER,
                     size_pt=_FONT_SIZE_BODY)

    # Основні відомості — зведений DIMS-аналіз
    lines = [
        f"Інтегральна оцінка DIMS: R_DIMS = {r_dims:.4f} → грейд {grade}.",
        "",
        "Індикатори:",
        f"  I_content  = {float(indicators.get('I_content', 0)):.4f}",
        f"  I_coord    = {float(indicators.get('I_coord', 0)):.4f}",
        f"  I_dynamics = {float(indicators.get('I_dynamics', 0)):.4f}",
        f"  I_impact   = {float(indicators.get('I_impact', 0)):.4f}",
        f"  I_source   = {float(indicators.get('I_source', 0)):.4f}",
    ]
    if flagged_pairs:
        lines.append("")
        lines.append(f"Підозрілих стилометричних пар: {len(flagged_pairs)}.")
        for fp in flagged_pairs[:5]:  # показуємо до 5
            a = (fp.get("a_source") or {}).get("display_title", "A")
            b = (fp.get("b_source") or {}).get("display_title", "B")
            d = fp.get("delta", 0)
            lines.append(f"  • {a} ↔ {b} (Δ = {float(d):.4f})")
    _clear_and_write(cells[3], "\n".join(lines), size_pt=_FONT_SIZE_BODY)

    # Пропозиції
    _clear_and_write(
        cells[4],
        _GRADE_RECOMMENDATIONS.get(grade, ""),
        size_pt=_FONT_SIZE_BODY,
    )


# ─────────────────────────────────────────────────────────────────────────────
#  ГОЛОВНА ФУНКЦІЯ
# ─────────────────────────────────────────────────────────────────────────────

def build_daily_monitoring_form(
    *,
    grade_info: dict,
    r_dims: float,
    manifestation: Optional[dict],
    indicators: dict,
    sources: list[dict],
    source_breakdown: Optional[list[dict]] = None,
    flagged_pairs: Optional[list[dict]] = None,
    direction: str = "",
    organization: str = "УПРАВЛІННЯ ЗАБЕЗПЕЧЕННЯ РЕАГУВАННЯ НА КРИЗОВІ СИТУАЦІЇ",
    custom_recommendation: str = "",
    compiled_at: Optional[datetime] = None,
) -> BytesIO:
    """Формує Анкету добового моніторингу (Додаток 1) у вигляді ``.docx``.

    Структура відтворює оригінальний бланк СЦМОУ:
    - горизонтальна таблиця з колонками відповідно до Методики;
    - секційні заголовки ІНФОРМАЦІЙНА СФЕРА / СТИЛОМЕТРИЧНИЙ АНАЛІЗ DIMS;
    - рядок-рядок для кожного опрацьованого джерела;
    - зведений рядок DIMS-індикаторів;
    - підписний блок.

    Повертає :class:`io.BytesIO` з готовим вмістом .docx.
    """
    compiled_at = compiled_at or datetime.now()
    grade = str(grade_info.get("grade") or "F").upper()
    manifestation = manifestation or {}
    manifestation_label = str(manifestation.get("label") or "")
    flagged_pairs = flagged_pairs or []
    source_breakdown = source_breakdown or []

    # Мітки джерел, що входять до підозрілих пар
    flagged_labels: set[str] = set()
    for fp in flagged_pairs:
        for key in ("a_source", "b_source"):
            src = fp.get(key) or {}
            lbl = src.get("label") or src.get("alias") or ""
            if lbl:
                flagged_labels.add(lbl)

    # Карта I_source за міткою джерела
    breakdown_map: dict[str, float] = {}
    for row in source_breakdown:
        lbl = (row.get("label") or
               (row.get("source") or {}).get("label") or "")
        score = row.get("score")
        if lbl and score is not None:
            breakdown_map[lbl] = float(score)

    # ── Документ ──────────────────────────────────────────────────────────
    doc = Document()
    for section in doc.sections:
        section.left_margin  = Cm(2.5)
        section.right_margin = Cm(1.5)
        section.top_margin   = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.page_width   = Cm(29.7)  # A4 альбом для широкої таблиці
        section.page_height  = Cm(21.0)

    # ── Шапка ─────────────────────────────────────────────────────────────
    # Організаційний заголовок (по центру, великими, жирний)
    p_org = doc.add_paragraph()
    p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_org = p_org.add_run(organization.upper())
    r_org.bold = True
    r_org.font.size = Pt(12)

    doc.add_paragraph()  # порожній рядок

    # Назва бланку
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t = p_title.add_run("Анкета добового моніторингу")
    r_t.bold = True
    r_t.font.size = Pt(12)

    # Підзаголовок (напрям + дата)
    subtitle_parts = []
    if direction:
        subtitle_parts.append(direction)
    subtitle_parts.append(
        f"станом на {compiled_at.strftime('%d.%m.%Y %H:%M')}"
    )
    p_sub = doc.add_paragraph(" ".join(subtitle_parts))
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.runs[0].font.size = Pt(11)

    doc.add_paragraph()

    # ── Таблиця ───────────────────────────────────────────────────────────
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Заголовки колонок
    hdr = table.rows[0].cells
    for idx, w in enumerate(_COL_WIDTHS):
        hdr[idx].width = w
        hdr[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_bg(hdr[idx], "D9D9D9")

    _col_titles = [
        "№",
        "Необхідність реагування\n(так/ні)",
        "Пріоритет\n(F/B/S/SS/SSS)",
        "Основні відомості",
        "Пропозиції щодо реагування / посилання",
    ]
    for idx, title in enumerate(_col_titles):
        _clear_and_write(
            hdr[idx], title,
            bold=True,
            size_pt=_FONT_SIZE_BODY,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    # ── Секція: ІНФОРМАЦІЙНА СФЕРА — по одному рядку на джерело ──────────
    _add_section_header(table, "ІНФОРМАЦІЙНА СФЕРА")

    for i, src in enumerate(sources, start=1):
        _add_source_row(
            table,
            num=i,
            source=src,
            grade=grade,
            manifestation_label=manifestation_label,
            custom_recommendation=custom_recommendation,
            flagged_labels=flagged_labels,
            breakdown_map=breakdown_map,
        )

    # ── Секція: СТИЛОМЕТРИЧНИЙ АНАЛІЗ DIMS ───────────────────────────────
    _add_section_header(table, "СТИЛОМЕТРИЧНИЙ АНАЛІЗ DIMS")

    _add_dims_summary_row(
        table,
        num=len(sources) + 1,
        r_dims=r_dims,
        grade=grade,
        indicators=indicators,
        flagged_pairs=flagged_pairs,
    )

    # Встановлення ширин усіх рядків (повторна гарантія)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(_COL_WIDTHS):
                cell.width = _COL_WIDTHS[idx]

    # ── Підписний блок ────────────────────────────────────────────────────
    doc.add_paragraph()
    p_sign = doc.add_paragraph()
    p_sign.add_run(
        "Начальник підрозділу моніторингу\n"
        "\n"
        "_________________________________    "
        "___________________________\n"
        "     (посада, звання)                           "
        "(підпис, ПІБ)"
    ).font.size = Pt(10)

    # Технічна примітка
    doc.add_paragraph()
    p_note = doc.add_paragraph()
    note_run = p_note.add_run(
        "Анкету сформовано автоматично програмним комплексом "
        "стилометричного моніторингу (DIMS) на виконання Методики щодо "
        "процедур моніторингу інформації у відкритих джерелах та її "
        "обробки (затв. наказом НУЗРКС МОУ № 46 від 28.11.2022). "
        f"Дата формування: {compiled_at.strftime('%d.%m.%Y %H:%M')}."
    )
    note_run.italic = True
    note_run.font.size = Pt(8)

    # ── Збереження ────────────────────────────────────────────────────────
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
